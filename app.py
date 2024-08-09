import os
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from docxcompose.composer import Composer
from pdf2docx import Converter
import io
import boto3
from botocore.exceptions import ClientError
from PIL import Image
import tempfile
import logging
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import json
import re

app = Flask(__name__)

AWS_ACCESS_KEY_ID = os.environ.get('AWS_ACCESS_KEY_ID')
AWS_SECRET_ACCESS_KEY = os.environ.get('AWS_SECRET_ACCESS_KEY')
S3_BUCKET_NAME = os.environ.get('S3_BUCKET_NAME')

if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
    print("Warning: One or more required environment variables are missing.")

# S3 configuration
s3_client = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY
)
s3_bucket_name=S3_BUCKET_NAME

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def is_image(file_content):
    try:
        Image.open(io.BytesIO(file_content))
        return True
    except IOError:
        return False

def is_pdf(filename):
    return filename.lower().endswith('.pdf')

def convert_pdf_to_docx(pdf_content):
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
            pdf_temp.write(pdf_content)
            pdf_temp_path = pdf_temp.name

        docx_temp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        docx_temp_path = docx_temp.name
        docx_temp.close()

        # Convert PDF to DOCX
        cv = Converter(pdf_temp_path)
        cv.convert(docx_temp_path)
        cv.close()

        # Read the converted DOCX
        with open(docx_temp_path, 'rb') as docx_file:
            docx_content = docx_file.read()

        # Clean up temporary files
        os.unlink(pdf_temp_path)
        os.unlink(docx_temp_path)

        return docx_content
    except Exception as e:
        logger.error(f"Error converting PDF to DOCX: {str(e)}")
        return None

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def process_files(files, template_name, date):
    merged_doc = Document()
    composer = Composer(merged_doc)
    
    file_types = request.form.getlist('file_types')
    captions = request.form.getlist('captions')
    
    logger.info(f"Processing {len(files)} files")
    
    for i, (file, file_type, caption) in enumerate(zip(files, file_types, captions + [None] * len(files))):
        try:
            logger.info(f"Processing file {i+1}: {file.filename}, Type: {file_type}")
            
            # If it's the invoice file (which is stored in S3), download it first
            if file.filename.startswith('invoice_'):
                s3_object = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=file.filename)
                file_content = s3_object['Body'].read()
            else:
                file_content = file.read()
            
            if file_type == 'Foto':
                logger.info(f"{file.filename} is an image")
                if caption:
                    para = merged_doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(caption)
                    run.bold = True
                    run.font.size = Pt(12)
                pil_image = Image.open(io.BytesIO(file_content))
                img_stream = io.BytesIO()
                pil_image.save(img_stream, format=pil_image.format)
                img_stream.seek(0)
                merged_doc.add_picture(img_stream, width=Inches(6))
            elif file_type == 'PDF':
                logger.info(f"{file.filename} is a PDF")
                docx_content = convert_pdf_to_docx(file_content)
                if docx_content is None:
                    logger.error(f"PDF conversion failed for {file.filename}")
                    continue
                doc = Document(io.BytesIO(docx_content))
                composer.append(doc)
            else:
                logger.info(f"{file.filename} is a document")
                doc = Document(io.BytesIO(file_content))
                composer.append(doc)
        except Exception as e:
            logger.error(f"Error processing {file.filename}: {str(e)}")
            logger.exception("Full traceback:")
            continue
    
    output_filename = f'SPJ_{template_name}_{date}.docx'
    output_stream = io.BytesIO()
    composer.save(output_stream)
    output_stream.seek(0)
    
    # Upload to S3
    try:
        s3_client.upload_fileobj(output_stream, S3_BUCKET_NAME, output_filename)
        logger.info(f"Uploaded merged file to S3: {output_filename}")
    except Exception as e:
        logger.error(f"Error uploading to S3: {str(e)}")
        raise
    
    return output_filename

def generate_invoice(data):
    try:
        json_data = json.dumps(data)
        result = subprocess.run(['/app/.heroku/node/bin/node', 'generate_invoice.js', json_data], capture_output=True, text=True, check=True)
        
        # Extract the output path from the Node.js script output
        output_path_match = re.search(r'OUTPUT_PATH:(.+)', result.stdout)
        if output_path_match:
            return output_path_match.group(1).strip()
        else:
            raise Exception("Failed to extract output path from Node.js script")
    except subprocess.CalledProcessError as e:
        print(f"Error running Node.js script: {e}")
        print(f"Script error output: {e.stderr}")
        raise Exception("Failed to generate invoice")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_spj():
    template_name = request.form.get('templateName')
    date = request.form.get('tanggalAcara')
    files = request.files.getlist('files')
    invoice_data = json.loads(request.form.get('invoiceData'))
    
    logger.info(f"Generating SPJ - Template: {template_name}, Date: {date}")
    logger.info(f"Number of files received: {len(files)}")
    
    if not template_name or not date or not files:
        return jsonify({'error': 'Missing required data'}), 400
    
    valid_files = [f for f in files if f and allowed_file(f.filename)]
    logger.info(f"Number of valid files: {len(valid_files)}")
    
    if not valid_files:
        return jsonify({'error': 'No valid files to process'}), 400
    
    try:
        # Generate invoice using the data from the frontend
        invoice_path = generate_invoice(invoice_data)
        logger.info(f"Invoice generated at: {invoice_path}")
        
        # Add the generated invoice to the list of files to process
        invoice_filename = os.path.basename(invoice_path)
        invoice_file_obj = type('', (), {})()
        invoice_file_obj.filename = invoice_filename
        valid_files.append(invoice_file_obj)
        file_types.append('Dokumen')  # Assuming invoice is always a document
        
        logger.info(f"Total files to process (including invoice): {len(valid_files)}")
        
        merged_filename = process_files(valid_files, template_name, date)
        logger.info(f"Merged file created: {merged_filename}")
        
        return jsonify({'success': True, 'file': merged_filename})
    except Exception as e:
        logger.error(f"Error processing files: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_obj = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=filename)
        return send_file(
            io.BytesIO(file_obj['Body'].read()),
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except ClientError as e:
        return str(e), 404

@app.route('/generate_kuitansi', methods=['POST'])
def generate_kuitansi():
    invoice_data = request.json
    logger.info(f"Received invoice data: {invoice_data}")
    
    try:
        # Generate invoice using the data from the frontend
        invoice_path = generate_invoice(invoice_data)
        logger.info(f"Invoice generated at: {invoice_path}")
        
        # Get the filename
        filename = os.path.basename(invoice_path)
        
        # Upload the file to S3
        with open(invoice_path, 'rb') as invoice_file:
            s3_client.upload_fileobj(invoice_file, S3_BUCKET_NAME, filename)
        logger.info(f"Invoice uploaded to S3: {filename}")
        
        # Remove the local file after uploading to S3
        os.remove(invoice_path)
        
        return jsonify({'success': True, 'filename': filename})
    except Exception as e:
        logger.error(f"Error generating kuitansi: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=False)